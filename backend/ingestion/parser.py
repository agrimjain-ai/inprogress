import os
import logging
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from dataclasses import dataclass
from typing import Optional, List, Dict, Any

# Configure structured production logging
logger = logging.getLogger(__name__)


# ==============================================================================
# PRODUCTION NOTE ON LICENSING:
# PyMuPDF (fitz) is licensed under AGPL v3 / Commercial. 
# If your production environment has strict licensing compliance (e.g., closed-source
# proprietary SaaS), you may need to replace PyMuPDF with an alternative library 
# like pypdf, pdfminer.six, or pdfplumber (MIT/Apache licensed).
# ==============================================================================

@dataclass(frozen=True)
class PageData:
    page_number: int
    text: str

@dataclass
class ParsedDocument:
    filename: str
    doc_type: str
    raw_text: str
    pages: List[PageData]
    total_pages: int


class DocumentParserError(Exception):
    """Base exception for parser errors."""
    pass


def detect_doc_type(text: str) -> str:
    """
    Uses a basic scoring mechanism to avoid first-match-wins bias.
    
    TODO FOR PRODUCTION:
    - This rule-based classification can be fragile and prone to false positives.
    - For high-accuracy classification, consider replacing this with a lightweight 
      machine learning classifier (e.g., TF-IDF + Logistic Regression) or a 
      structured prompt utilizing an LLM API.
    """
    text_lower = text.lower()
    
    scores = {
        "judgment": sum(word in text_lower for word in ["honourable", "appellant", "respondent", "vs", "versus", "tribunal", "court", "judgment"]),
        "circular": sum(word in text_lower for word in ["circular no.", "cbic", "clarification", "circular", "trade notice"]),
        "notification": sum(word in text_lower for word in ["notification no.", "s.o.", "g.s.r", "notification", "gazette"]),
        "act": sum(word in text_lower for word in ["section", "chapter", "act, 2017", "cgst act", "igst act"])
    }
    
    best_match, highest_score = max(scores.items(), key=lambda x: x[1])
    return best_match if highest_score > 0 else "unknown"


def parse_pdf(file_path: str, filename: str) -> ParsedDocument:
    """
    Parses digital PDF files.
    
    TODO FOR PRODUCTION (OCR & CLOUD STORAGE):
    1. OCR (Scanned PDFs): This code only extracts digital/selectable text. If a PDF 
       is scanned (just images), `page.get_text()` will return empty. You will need to 
       integrate an OCR engine (e.g., Tesseract via pytesseract, pdfplumber, or cloud APIs 
       like AWS Textract/Azure Document Intelligence) as a fallback when text is empty.
    2. Cloud Storage Integration: Instead of taking local `file_path`, a production pipeline 
       often receives files as binary streams or bytes from an S3 bucket or API request. 
       You will need to adapt this to open from memory streams (e.g., using `io.BytesIO`).
    """
    logger.info("Starting PDF parsing: %s", filename)
    pages: List[PageData] = []
    full_text_parts: List[str] = []

    try:
        # Utilizing context manager for defensive resource management
        with fitz.open(file_path) as doc:
            if doc.is_encrypted:
                raise DocumentParserError(f"PDF is password-protected or encrypted: {filename}")

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = (page.get_text() or "").strip()

                if text:
                    pages.append(PageData(page_number=page_num + 1, text=text))
                    full_text_parts.append(text)
            
            total_pages = len(doc)
            
    except Exception as e:
        logger.exception("Failed to parse PDF: %s", filename)
        raise DocumentParserError(f"PDF parsing failed: {e}") from e

    full_text = "\n".join(full_text_parts)
    return ParsedDocument(
        filename=filename,
        doc_type=detect_doc_type(full_text),
        raw_text=full_text,
        pages=pages,
        total_pages=total_pages
    )


def parse_docx(file_path: str, filename: str) -> ParsedDocument:
    """
    Parses Word documents.
    
    TODO FOR PRODUCTION:
    1. Media/Images: This code ignores embedded images. If those images contain crucial 
       information, you would need to extract and OCR them.
    2. True Pagination: Word files do not have standard, fixed pages in their XML structure. 
       The simulated pagination below is a fallback. If your application strictly relies 
       on exact matching to printed pages, you may need a heavy layout engine or PDF-conversion 
       step to compute true pages.
    """
    logger.info("Starting DOCX parsing: %s", filename)
    pages: List[PageData] = []
    full_text_parts: List[str] = []
    
    try:
        doc = DocxDocument(file_path)
        
        # Extract paragraph text
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text_parts.append(text)

        # Extract table text (critical for business/legal documents)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text_parts.append(" | ".join(row_text))

    except Exception as e:
        logger.exception("Failed to parse DOCX: %s", filename)
        raise DocumentParserError(f"DOCX parsing failed: {e}") from e

    full_text = "\n".join(full_text_parts)
    
    # Simulating pages via character limits (approx. 1-2 pages of text)
    chunk_size = 3000
    chunks = [full_text[i:i + chunk_size].strip() for i in range(0, len(full_text), chunk_size)]
    
    pages = [
        PageData(page_number=idx + 1, text=chunk)
        for idx, chunk in enumerate(chunks) if chunk
    ]

    return ParsedDocument(
        filename=filename,
        doc_type=detect_doc_type(full_text),
        raw_text=full_text,
        pages=pages,
        total_pages=max(len(pages), 1)
    )


def parse_document(file_path: str, filename: str) -> ParsedDocument:
    """
    Main entrypoint for document parsing.
    
    TODO FOR PRODUCTION (SECURITY):
    - Do not trust file extensions. An attacker can upload a malicious script renamed 
      with a '.pdf' extension. Use file signature verification (e.g., using python-magic 
      to read magic bytes) to verify the actual file MIME type before processing.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    ext = os.path.splitext(filename.lower())[1]
    
    if ext == ".pdf":
        return parse_pdf(file_path, filename)
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_path, filename)
    else:
        raise ValueError(f"Unsupported file extension '{ext}' for file: {filename}")